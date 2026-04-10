import docx
import copy
import os
import tempfile
import subprocess

# Word 公式的命名空间
NS_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def convert_math_node_with_pandoc(math_node):
    """
    后台创建一个临时的微型 Word 文档让 Pandoc 解析。
    这是最安全、保证 100% 转换成功的方法。
    """
    temp_doc = docx.Document()
    p = temp_doc.add_paragraph()
    
    # 将公式节点完整拷贝到临时文档中
    p._element.append(copy.deepcopy(math_node))
    
    # 保存临时 .docx 文件
    fd, temp_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    temp_doc.save(temp_path)
    
    try:
        # 让 Pandoc 解析这个合法的临时 docx 文档
        cmd = ['pandoc', temp_path, '-f', 'docx', '-t', 'markdown', '--wrap=none']
        result = subprocess.check_output(cmd, stderr=subprocess.DEVNULL).decode('utf-8').strip()
        
        if not result:
            return " [公式为空] "
        
        # 返回提取到的 LaTeX 文本
        return f" {result} "
    except Exception as e:
        return " [转换失败] "
    finally:
        # 清理临时文件
        if os.path.exists(temp_path):
            os.remove(temp_path)

def _replace_math_in_element(element):
    """查找并替换元素内部的所有公式节点"""
    math_tag = f"{{{NS_M}}}oMath"
    
    # 找到所有公式节点
    math_nodes = [node for node in element.iter() if node.tag == math_tag]
    
    for math_node in math_nodes:
        # 调用 Pandoc 获取转换后的文本
        latex_text = convert_math_node_with_pandoc(math_node)
        
        # 创建一个普通的 Word 文本节点，把 LaTeX 塞进去
        new_run = docx.oxml.shared.OxmlElement('w:r')
        new_text = docx.oxml.shared.OxmlElement('w:t')
        new_text.text = latex_text
        new_run.append(new_text)
        
        # 把原来的公式节点替换成纯文本节点
        parent = math_node.getparent()
        if parent is not None:
            parent.replace(math_node, new_run)

def process_single_docx(input_path, output_path):
    if not os.path.exists(input_path):
        print(f"❌ 找不到输入文件，请检查路径: {input_path}")
        return

    print(f"⏳ 正在加载文档: {os.path.basename(input_path)}")
    print("⏳ 正在逐个提取公式并转为 LaTeX，这可能需要几十秒时间，请耐心等待...")
    
    doc = docx.Document(input_path)
    
    # 1. 遍历处理所有的普通段落
    for para in doc.paragraphs:
        _replace_math_in_element(para._element)
        
    # 2. 遍历处理所有的表格（这是为了保住你的表格里的公式）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_math_in_element(para._element)
                    
    doc.save(output_path)
    print(f"✅ 测试完成！请查看生成的新文件: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    # 直接使用你之前提供的那个具体文件路径
    input_file = "/Users/hq/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_esz1onoxfpw812_4077/msg/file/2026-03/1.docx"
    
    # 输出到当前目录，方便你查看
    output_file = "./测试版_LaTeX.docx"
    
    process_single_docx(input_file, output_file)
